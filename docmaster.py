# Author: Adrian McPherson
# Contact: adrian[dot]mcpherson[at]gmail[dot][com]

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from docx2pdf import convert
from PyPDF2 import PdfMerger
from docx.enum.section import WD_SECTION_START


def create_template(class_code, subject, teacher_name, term, unit_number, unit_title, assignment_name):
    template = docx.Document()

    # Set up header
    header = template.sections[0].header
    header_paragraph = header.paragraphs[0]

    # Add class code, subject, and teacher name to the left side of the header
    if class_code or subject or teacher_name:
        left_header_run = header_paragraph.add_run(f"{class_code} - {subject} - {teacher_name}".strip("- "))
        left_header_run.font.size = Pt(10)
        left_header_run.font.name = "Calibri"

        # Add line break
        header_paragraph.add_run("\n")

    # Add term, unit number, unit title, and assignment name to the left side of the header
    term_string = f"Term {term}" if term != "Select" else ""
    unit_string = f"Unit {unit_number}:" if unit_number else ""
    header_info = f"{term_string} {unit_string} {unit_title} - {assignment_name}".strip("- ,:")
    if header_info:
        left_header_run = header_paragraph.add_run(header_info)
        left_header_run.font.size = Pt(10)
        left_header_run.font.name = "Calibri"

    return template


def browse_files():
    global files
    files = filedialog.askopenfilenames(filetypes=[("Word Documents", "*.docx")])
    selected_files_label.config(text=f"{len(files)} files selected")
    selected_files_label = tk.Label(root, text="")
    selected_files_label.grid(row=11, column=0, columnspan=2, padx=10, pady=10, sticky=tk.W + tk.E)



def process_files():
    class_code = e1.get()
    subject = e2.get()
    teacher_name = e3.get()
    term = term_var.get()
    unit_number = e4.get()
    unit_title = e5.get()
    assignment_name = e6.get()
    include_filename = include_filename_var.get()
    double_spacing = double_spacing_var.get()
    standardize_font = standardize_font_var.get()
    create_pdf = create_pdf_var.get()

    template = create_template(class_code, subject, teacher_name, term, unit_number, unit_title, assignment_name)

    pdf_files = []
    
    add_page_numbers = add_page_numbers_var.get()
    
    for file in files:
        doc = docx.Document(file)

        # Copy header content from template to document
        for i, paragraph in enumerate(template.sections[0].header.paragraphs):
            if len(doc.sections[0].header.paragraphs) <= i:
                doc.sections[0].header.add_paragraph()
            target_paragraph = doc.sections[0].header.paragraphs[i]
            target_paragraph.clear()
            for run in paragraph.runs:
                target_run = target_paragraph.add_run(run.text)
                target_run.font.size = run.font.size
                target_run.font.name = run.font.name

        if include_filename:
            # Add file name without .docx extension and a line break to the header
            file_name = file.split("/")[-1].replace(".docx", "")
            header_run = doc.sections[0].header.paragraphs[0].add_run(f"\n{file_name}\n")
            header_run.font.size = Pt(10)
            header_run.font.name = "Calibri"

        # Add footer with "Page X of Y" format if the checkbox is selected
        if add_page_numbers:
            footer = doc.sections[0].footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add "Page " text
            footer_run = footer_paragraph.add_run("Page ")
            footer_run.font.size = Pt(10)
            footer_run.font.name = "Calibri"

    # Add the current page number field
            page_num_run = footer_paragraph.add_run()
            page_num_run._element.append(OxmlElement('w:fldSimple'))
            page_num_run._element[-1].set(qn('w:instr'), 'PAGE')

    # Add " of " text
            footer_run = footer_paragraph.add_run(" of ")
            footer_run.font.size = Pt(10)
            footer_run.font.name = "Calibri"

    # Add the total number of pages field
            total_pages_run = footer_paragraph.add_run()
            total_pages_run._element.append(OxmlElement('w:fldSimple'))
            total_pages_run._element[-1].set(qn('w:instr'), 'NUMPAGES')

        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if double_spacing:
                paragraph.paragraph_format.line_spacing = 2.0

            if standardize_font:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = "Calibri"

        doc.save(file)

        if create_pdf:
            pdf_file = file.replace(".docx", ".pdf")
            convert(file, pdf_file)
            pdf_files.append(pdf_file)

    if create_pdf and len(pdf_files) > 0:
        # Merge PDFs
        merger = PdfMerger()
        for pdf_file in pdf_files:
            with open(pdf_file, "rb") as f:
                merger.append(f)

        first_file_directory = os.path.dirname(files[0])
        merged_output = os.path.join(first_file_directory, "merged_output.pdf")
        with open(merged_output, "wb") as output_file:
            merger.write(output_file)

        # Close and delete individual PDFs
        for pdf_file in pdf_files:
            if pdf_file != merged_output:
                os.remove(pdf_file)

    messagebox.showinfo("Success", f"{len(files)} files processed successfully.")


root = tk.Tk()
root.title("DOC Master")

# Labels and entries
tk.Label(root, text="Class Code:").grid(row=0, column=0, padx=10, pady=10, sticky=tk.E)
tk.Label(root, text="Subject:").grid(row=1, column=0, padx=10, pady=10, sticky=tk.E)
tk.Label(root, text="Teacher Name:").grid(row=2, column=0, padx=10, pady=10, sticky=tk.E)
tk.Label(root, text="Term:").grid(row=3, column=0, padx=10, pady=10, sticky=tk.E)
tk.Label(root, text="Unit Number:").grid(row=4, column=0, padx=10, pady=10, sticky=tk.E)
tk.Label(root, text="Unit Title:").grid(row=5, column=0, padx=10, pady=10, sticky=tk.E)
tk.Label(root, text="Assignment Name:").grid(row=6, column=0, padx=10, pady=10, sticky=tk.E)

e1 = tk.Entry(root)
e2 = tk.Entry(root)
e3 = tk.Entry(root)
term_var = tk.StringVar()
term_dropdown = ttk.Combobox(root, textvariable=term_var, values=["Select"] + list(range(1, 5)), state="readonly")
term_dropdown.current(0)
e4 = tk.Entry(root)
e5 = tk.Entry(root)
e6 = tk.Entry(root)

e1.grid(row=0, column=1, padx=10, pady=10)
e2.grid(row=1, column=1, padx=10, pady=10)
e3.grid(row=2, column=1, padx=10, pady=10)
term_dropdown.grid(row=3, column=1, padx=10, pady=10)
e4.grid(row=4, column=1, padx=10, pady=10)
e5.grid(row=5, column=1, padx=10, pady=10)
e6.grid(row=6, column=1, padx=10, pady=10)

# Checkboxes
standardize_font_var = tk.BooleanVar()
standardize_font_checkbox = tk.Checkbutton(root, text="Standardize font", variable=standardize_font_var)
standardize_font_checkbox.grid(row=7, column=0, padx=10, pady=10, sticky=tk.W)

double_spacing_var = tk.BooleanVar()
double_spacing_checkbox = tk.Checkbutton(root, text="Double spacing", variable=double_spacing_var)
double_spacing_checkbox.grid(row=7, column=1, padx=10, pady=10, sticky=tk.W)

include_filename_var = tk.BooleanVar()
include_filename_checkbox = tk.Checkbutton(root, text="Include filename in header", variable=include_filename_var)
include_filename_checkbox.grid(row=8, column=0, padx=10, pady=10, sticky=tk.W)

add_page_numbers_var = tk.BooleanVar()
add_page_numbers_checkbox = tk.Checkbutton(root, text="Add page numbers", variable=add_page_numbers_var)
add_page_numbers_checkbox.grid(row=9, column=0, padx=10, pady=10, sticky=tk.W)

create_pdf_var = tk.BooleanVar()
create_pdf_checkbox = tk.Checkbutton(root, text="Create merged PDF", variable=create_pdf_var)
create_pdf_checkbox.grid(row=8, column=1, padx=10, pady=10, sticky=tk.W)

# Buttons
browse_button = ttk.Button(root, text="Browse Files", command=browse_files)
browse_button.grid(row=10, column=0, padx=10, pady=10, sticky=tk.W+tk.E)

process_button = ttk.Button(root, text="Process Files", command=process_files)
process_button.grid(row=10, column=1, padx=10, pady=10, sticky=tk.W+tk.E)

root.mainloop()

